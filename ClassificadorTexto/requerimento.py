# -*- coding: utf-8 -*-
"""
@author: Giliard
"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
import os
import re

def obterTexto(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


def limparTextoPorEsquema(documento,esquema={}):
    documento = [ line.strip() for line in documento ]
    documento = [ (line) for line in documento if len(line) is not 0 ]
    documento = '\n'.join(documento)

    for expressao in esquema.keys():
        documento = re.sub(expressao,esquema[expressao], documento)

    return documento

def relacionarRequerimentos(documento,documentoRelacaoRequerimentos):

    documento = documento.split('\n')
    
    maxIndice = len(documento)
    indice = 0
    # relacao = list()
    
    while indice < maxIndice:
        linha = documento[indice]
        if linha == 'BEGIN_DOCUMENT':
            final = indice + 1
            tipo = documento[indice-1]
            texto = list()
            while final < maxIndice:
                linha = documento[final]
                if linha == 'END_DOCUMENT':
                    break
                texto.append(linha + ' ')
                final += 1
            indice = final + 1

            texto = ''.join(texto)
            textoRequerimento = ', de autoria d' + texto[0].lower() + texto[1:]
            # requerimento = tipo + textoRequerimento
            # relacao.append(requerimento)
            documentarRelacaoRequerimentos(tipo,textoRequerimento,documento=documentoRelacaoRequerimentos)
        else :
            indice += 1
    

def documentarRelacaoRequerimentos(tipoRequerimento,textoRequerimento,documento):
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = ALIGN.JUSTIFY
    
    run = paragrafo.add_run(tipoRequerimento)
    run.bold = True
    run.font.name = "Times New Romam"
    run.font.size = Pt(12)

    run = paragrafo.add_run(textoRequerimento)
    run.font.name = "Times New Romam"
    run.font.size = Pt(12)


if __name__ == "__main__":

    limpar = {
        'Excelentíssimo Senhor Odemir Jacob' : '',
        'Presidente da Câmara Municipal de Santo Antônio da Platina' : 'BEGIN_DOCUMENT',
        'Nestes termos,' : '',
        'Pede deferimento.' : '',
        'infra-assinad[ao],' : '',
        'no uso das atribuições que lhes são conferidas pelo Regimento Interno,' : '',
        'requer a Vossa Excelência que seja expedido ofício': '',
        'SECRETARIA DA CÂMARA MUNICIPAL DE SANTO ANTÔNIO DA PLATINA, ESTADO DO PARANÁ,' : '\nEND_DOCUMENT\n',
        'em \d+ de \w+ de \d+' : '',
        '\n\n+' : '\n'
    }   

    diretorio = os.path.join('data','15')
    arquivos = os.listdir(diretorio)

    # requerimentos = list()
    saida = docx.Document()
    saida.core_properties.language = "pt-BR"

    for word_file in  arquivos:
        word_text = obterTexto(os.path.join(diretorio,word_file))
        word_text = limparTextoPorEsquema(word_text,limpar)
        relacionarRequerimentos(word_text,documentoRelacaoRequerimentos=saida)
    
    saida.save('requerimentos.docx')